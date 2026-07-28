# -*- coding: utf-8 -*-
"""Genera la carpeta 'escritorio-del-pasante': el entorno introductorio a Cowork
de la Misión 0 (universo Atlas/Faro). Regenerable: borrar la carpeta y volver a correr.

Requiere: openpyxl, python-docx, reportlab.
Filosofía (adaptada de caso_finanzas_desordenado, DIAA2026): carpeta chica pero
completa, 4 formatos, archivos que se referencian entre sí, 4 errores plantados
verificables cruzando documentos, y regenerable por script.
"""
import os
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

BASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "escritorio-del-pasante")
os.makedirs(BASE, exist_ok=True)

# Datos "reales" de la semana (ground truth):
# YPF   cierre 42,10  anterior 40,85  -> variación +3,1%
# VIST  cierre 51,30  anterior 50,75  -> variación +1,1%
# TS    cierre 33,60  anterior 33,88  -> variación -0,8%
# Volumen operado (millones de acciones): YPF 12,4 · VIST 3,1 · TS 5,2 -> TOTAL 20,7

# ---------- 1) Excel: cotizaciones semana (1).xlsx ----------
wb = Workbook()
ws = wb.active
ws.title = "Cotizaciones"
ws.append(["Atlas S.A. - Seguimiento semanal de mercado - semana del 13 al 17 de julio de 2026 (USD)"])
ws.append([])
ws.append(["Empresa", "Ticker", "Cierre viernes", "Cierre semana anterior", "Variación %"])
ws.append(["YPF", "YPF", 42.10, 40.85, "+1,2%"])   # ERROR plantado E1: la variación real es +3,1%
ws.append(["Vista Energy", "VIST", 51.30, 50.75, "+1,1%"])
ws.append(["Tenaris", "TS", 33.60, 33.88, "-0,8%"])
ws2 = wb.create_sheet("Volumen")
ws2.append(["Volumen operado en la semana (millones de acciones)"])
ws2.append([])
ws2.append(["Ticker", "Volumen"])
ws2.append(["YPF", 12.4])
ws2.append(["VIST", 3.1])
ws2.append(["TS", 5.2])
ws2.append(["TOTAL", 19.7])  # ERROR plantado E2: la suma real es 20,7
wb.save(os.path.join(BASE, "cotizaciones semana (1).xlsx"))

# ---------- 2) Word: Pulso semanal - FINAL final.docx ----------
doc = Document()
doc.add_heading("Pulso semanal de mercado - semana del 13 al 17 de julio", 0)
p = doc.add_paragraph("Atlas S.A. · Seguimiento de mercado para la reunión de los lunes")
p.runs[0].font.size = Pt(11)
doc.add_paragraph("Preparado por: el pasante (borrador)")

doc.add_heading("1. Resumen de la semana", 1)
doc.add_paragraph(
    "Semana positiva para el sector: las tres empresas que seguimos cerraron al alza, "  # ERROR plantado E3: Tenaris bajó (-0,8%, ver planilla)
    "empujadas por el anuncio de nuevos pozos en Vaca Muerta y la demanda de energía "
    "de los centros de datos. YPF fue la de mejor desempeño de la semana. Los cierres "
    "y variaciones están en la planilla `cotizaciones semana (1).xlsx`."
)
doc.add_heading("2. Noticias por empresa", 1)
doc.add_paragraph(
    "YPF: anunció la puesta en marcha de dos pozos nuevos en Vaca Muerta; la acción "
    "acompañó con la mejor suba del trío. Vista Energy: presenta resultados el jueves "
    "que viene; el mercado espera confirmación del guidance anual. Tenaris: semana sin "
    "anuncios propios; el movimiento respondió al precio internacional del acero."
)
doc.add_heading("3. Qué vigilar", 1)
doc.add_paragraph("[PENDIENTE DE COMPLETAR: los resultados de Vista del jueves y el dato de actividad de perforación]")
doc.save(os.path.join(BASE, "Pulso semanal - FINAL final.docx"))
# Nota: el borrador usa el formato viejo de 3 secciones; la guía vigente (PDF v3) exige 5.

# ---------- 3) PDF: guia-formato-pulso_v3.pdf ----------
styles = getSampleStyleSheet()
pdf = SimpleDocTemplate(os.path.join(BASE, "guia-formato-pulso_v3.pdf"), pagesize=A4,
                        topMargin=2.5 * cm, bottomMargin=2.5 * cm)
story = [
    Paragraph("Atlas S.A.", styles["Title"]),
    Paragraph("Guía de formato del pulso semanal - versión 3 (vigente desde julio 2026)", styles["Heading2"]),
    Spacer(1, 12),
    Paragraph(
        "Todo pulso semanal que llegue al equipo debe respetar la estructura de esta guía. "
        "Se envía los lunes a las 8:00, antes de la reunión de las 9:00.", styles["BodyText"]),
    Spacer(1, 12),
]
data = [
    ["Sección", "Contenido obligatorio"],
    ["1. Resumen de la semana", "Tres a cinco líneas, sin jerga financiera"],
    ["2. Tarjeta por empresa", "Cierre, variación y la noticia principal de cada una"],
    ["3. Tabla resumen", "Empresa · cierre · variación %, tomada de la planilla"],
    ["4. Qué vigilar", "Los eventos de la semana entrante"],
    ["5. Aclaración legal", "Uso interno; nunca es recomendación de inversión"],
]
t = Table(data, colWidths=[6 * cm, 9 * cm])
t.setStyle(TableStyle([
    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
]))
story += [t, Spacer(1, 16), Paragraph(
    "La aclaración legal es obligatoria en todos los envíos: los pulsos son síntesis de "
    "información pública para uso interno del equipo y nunca constituyen recomendaciones "
    "de inversión. Esta versión reemplaza a todas las anteriores.",
    styles["BodyText"])]
pdf.build(story)

# ---------- 4) Markdown: notas-del-pasante.md ----------
with open(os.path.join(BASE, "notas-del-pasante.md"), "w", encoding="utf-8") as f:
    f.write("""# notas de salida (último día del pasante)

lo que dejo a medias, perdón por el desorden

- la planilla `cotizaciones semana (1).xlsx` tiene los cierres del viernes YA cargados,
  pero calculé las variaciones a mano y la de YPF no me convence, revisarla
- el total de volumen tampoco me cierra, lo sumé apurado
- el borrador del pulso está en el Word "FINAL final" (sí, ya sé). OJO: creo que
  Tenaris terminó abajo esta semana, y el borrador dice otra cosa. verificar contra la planilla
- el formato que vale es el del PDF v3 (5 secciones, con tabla y aclaración legal).
  el TXT viejo que anda dando vueltas es de la época anterior, no usarlo
- falta la sección "qué vigilar" del pulso

pendientes
- [ ] verificar variación de YPF y el total de volumen
- [ ] corregir el resumen (¿las tres al alza?)
- [ ] pasar el borrador al formato de la guía v3
- [ ] renombrar estos archivos con algún criterio, esto es un caos
""")

# ---------- 5) TXT: Copia de notas viejas (recuperado).txt ----------
with open(os.path.join(BASE, "Copia de notas viejas (recuperado).txt"), "w", encoding="utf-8") as f:
    f.write("""notas formato pulso (VIEJO, ver la guia v3 en pdf)
el pulso se manda los viernes a ultima hora, 3 secciones: resumen / noticias / pendientes
sin aclaracion legal, eso lo agrega cada uno si quiere
esto quedo desactualizado cuando el jefe aprobo la guia v3 en julio.
conservar solo como referencia historica.
""")  # ERROR plantado E4: contradice a la guía vigente (día de envío, secciones, aclaración legal)

print("Escritorio del pasante generado en", BASE)
for n in sorted(os.listdir(BASE)):
    print(" -", n)
